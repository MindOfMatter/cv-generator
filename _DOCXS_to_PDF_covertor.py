import json
import subprocess
import os
import sys
from docx import Document
import shutil
import time

# Constants
# Constant for the setup configuration file path
SETUP_CONFIG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'JSON/setup.json')

def load_json_file(file_path):
    """Load JSON file from the given path."""
    with open(file_path, 'r', encoding='utf-8') as file:
        print(f"Loading JSON from {file_path}")
        return json.load(file)

# Load configuration from setup.json
config = load_json_file(SETUP_CONFIG_PATH)

# Use the configuration values
LIBRE_OFFICE_PATH = config['LIBRE_OFFICE_PATH']
TMP_MERGED_FILENAME = '.tmp_merged.docx'
PDF_EXTENSION = '.pdf'
DOCX_EXTENSION = '.docx'
BASE_PATH = os.path.dirname(os.path.abspath(__file__))
TMP_MERGED_DOCX = os.path.join(BASE_PATH, 'RESULTS', TMP_MERGED_FILENAME)

def merge_docx(file1, file2, merged_file):
    # Check if the first file exists
    if not os.path.exists(file1):
        # If the first file doesn't exist, just use the second file
        merged_document = Document(file2)
    else:
        # If the first file exists, start merging
        merged_document = Document(file1)

        # Add a page break if the first file is not empty
        if merged_document.paragraphs:
            merged_document.add_page_break()

        # Load the second document
        doc2 = Document(file2)
        for element in doc2.element.body:
            # Append each element of the second document to the merged document
            merged_document.element.body.append(element)

    # Save the merged document
    merged_document.save(merged_file)

def convert_to_pdf(source_file):
    """Convert a DOCX file to a PDF file using LibreOffice."""
    output_dir = os.path.dirname(source_file)
    print(f"Converting '{source_file}' to PDF in '{output_dir}'")

    subprocess.call([
        LIBRE_OFFICE_PATH,
        '--headless',
        '--norestore',
        '--nofirststartwizard',
        '--convert-to', 'pdf',
        source_file,
        '--outdir', output_dir
    ])

    print("Conversion initiated, waiting for completion...")
    time.sleep(3)  # Adjust the delay if necessary
  
def rename_and_move_pdf(source_file, output_file):
    """Rename and move the converted PDF file to the desired location."""
    # Construct the temporary and final output filenames
    tmp_pdf_output = os.path.splitext(source_file)[0] + '.pdf'
    final_pdf_output = os.path.splitext(output_file)[0] + '.pdf'

    # Check if the temporary converted file exists
    if os.path.exists(tmp_pdf_output):
        # If the final output file already exists, remove it
        if os.path.exists(final_pdf_output):
            print(f"Removing existing PDF document: {final_pdf_output}")
            os.remove(final_pdf_output)
        
        # Move the temporary file to the final output filename
        print(f"Updating the PDF document: {final_pdf_output}")
        shutil.move(tmp_pdf_output, final_pdf_output)
    else:
        print(f"Conversion failed or output file not found: {tmp_pdf_output}")

def remove_file(file_path):
    """Remove a file if it exists."""
    if os.path.exists(file_path):
        print(f"Removing the document: {file_path}")
        os.remove(file_path)

def finalize_to_pdf(source_file, output_file):
    convert_to_pdf(source_file)
    rename_and_move_pdf(source_file, output_file)
    remove_file(source_file)
    final_pdf_output = os.path.splitext(output_file)[0] + '.pdf'
    print("PDF Document Generated Successfully.")
    print(f"Opening the PDF document: {final_pdf_output}")
    os.startfile(final_pdf_output)
    
def is_valid_docx(file):
    return file.lower().endswith(DOCX_EXTENSION) and os.path.exists(file)

def main():
    """Main function to merge DOCX files and convert to PDF."""
    if len(sys.argv) > 1:
        docx_files = [f for f in sys.argv[1:] if is_valid_docx(f)]
    elif len(sys.argv) == 1:
        source_file = sys.argv[1]
        if is_valid_docx(source_file):
            convert_to_pdf(source_file)
            sys.exit(0)
    else:
        print("Missing parameters")
        sys.exit(1)

    if len(docx_files) < 2:
        print("At least two DOCX files are required.")
        sys.exit(1)

    for file in reversed(docx_files):
        merge_docx(TMP_MERGED_DOCX, file, TMP_MERGED_DOCX)

    lastFile = docx_files[-1]
    finalize_to_pdf(TMP_MERGED_DOCX, lastFile)

if __name__ == "__main__":
    main()
