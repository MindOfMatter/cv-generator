import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
import datetime
import subprocess
import os
import re
import sys

# Define margin values in inches
MARGIN = 0.75
LEFT_MARGIN = MARGIN
RIGHT_MARGIN = MARGIN
TOP_MARGIN = 0.5
BOTTOM_MARGIN = MARGIN

# Define font styles
HEADING_FONT_SIZE = Pt(14)
PARAGRAPH_FONT_SIZE = Pt(11)

# Constant for the setup configuration file path
SETUP_CONFIG_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'JSON/setup.json')

def load_json_file(file_path):
    """Load JSON file from the given path."""
    with open(file_path, 'r', encoding='utf-8') as file:
        print(f"Loading JSON from {file_path}")
        return json.load(file)

# Load configuration from setup.json
config = load_json_file(SETUP_CONFIG_PATH)

# Use the configuration values
LIBRE_OFFICE_PATH = config['LIBRE_OFFICE_PATH']
LANGUAGE = config['LANGUAGE']

def format_current_date():
    now = datetime.datetime.now()

    if LANGUAGE == "FR":
        # French months
        months = [
            "janvier", "février", "mars", "avril", "mai", "juin", 
            "juillet", "août", "septembre", "octobre", "novembre", "décembre"
        ]
        # Format the date - '29 mai 2017' format
        return f"{now.day} {months[now.month - 1]} {now.year}"
    else:  # English (Canada) date format
        # Format the date - 'May 29, 2017' format
        return now.strftime("%B %d, %Y")


def open_docx_with_libreoffice(file_path):
    """Open a DOCX file with LibreOffice, bypassing the recovery process."""
    try:
        subprocess.Popen([LIBRE_OFFICE_PATH, '--norestore', file_path])
    except Exception as e:
        print(f"An error occurred while opening the file with LibreOffice: {e}")

def open_docx(file_path):
    """Open a DOCX file with the default application or LibreOffice if available."""
    if os.path.exists(LIBRE_OFFICE_PATH):
        open_docx_with_libreoffice(file_path)
    else:
        try:
            os.startfile(file_path)
        except Exception as e:
            print(f"An error occurred while opening the file with the default application: {e}")

def merge_json_data(json_paths):
    merged_data = {}
    for key, path in json_paths.items():
        try:
            with open(path, 'r', encoding='utf-8') as file:
                print(f"Loading JSON from {path}")
                data = json.load(file)
                merged_data[key] = data
        except Exception as e:
            print(f"Error loading JSON file {path}: {e}")
    return merged_data

def get_data(data, key, current_item=None, indexes=None):
    """Retrieve data for a given key."""
    print(f"Attempting to retrieve data for key: {key}")  # Debug print

    if key == None:
        return None

    if key == 'item':
        return current_item
    
    if key == 'current_date':
        return format_current_date()

    try:
        for k in key.split('.'):
            if indexes is not None:
                k = replace_indexes(k, indexes)

            data = traverse_data(data, k)

        return data if data is not None else ''
    except (TypeError, KeyError) as e:
        print(f"Error accessing key '{key}': {e}")  # Debug print
        return ''  # Return an empty string if any error occurs
    
def hex_to_rgb(hex_color):
    """Convert a hex color to an RGB tuple."""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

def set_run_style(run, font_size=None, hex_color=None):
    """Set the style for a run."""
    if font_size:
        run.font.size = Pt(font_size)
    if hex_color:
        rgb_color = hex_to_rgb(hex_color)
        run.font.color.rgb = RGBColor(*rgb_color)

def replace_indexes(key, indexes):
    """Replace placeholders with actual index values."""
    return re.sub(r"\{(\d+)\}", lambda m: str(indexes[int(m.group(1))]) if int(m.group(1)) < len(indexes) else m.group(0), key)

def traverse_data(data, key):
    """Traverse data based on the key."""
    if key.isdigit() and isinstance(data, list):
        key = int(key)
        if key < len(data):
            print(f"Accessing list index: {key}")  # Debug print
            return data[key]
        else:
            print(f"Warning: Index '{key}' out of range for the list.")  # Debug print
            return ''
    elif isinstance(data, dict) and key in data:
        print(f"Accessing dictionary key: {key}")  # Debug print
        return data[key]
    else:
        print(f"Warning: Key or index '{key}' not found in the data.")  # Debug print
        return ''

def format_string_with_data(format_str, data_points):
    """Format a string with given data points."""
    try:
        print(f"Formatting '{format_str}' with {data_points}")
        formatted_str = format_str.format(*data_points)
        print(f"Formatted '{formatted_str}'")
        return formatted_str
    except IndexError as e:
        print(f"Error formatting string: '{format_str}' with data points: {data_points} - {e}")  # Debug print
        return "Format Error"

def process_content(content, data, current_item=None, indexes=[]):
    """Process content based on the provided content definitions and data."""
    processed = []
    for item in content:
        keys = item.get('keys', [])
        format_str = item.get('format', '{0}')
        data_points = [get_data(data, key, current_item, indexes) for key in keys]
        processed.append(format_string_with_data(format_str, data_points))
    return processed

def process_nested_items(nested_each_data, item, data, indexes):
    """Process nested items within a data structure."""
    nested_items = get_data(item, nested_each_data['key'], None, indexes)
    print(f"Nested items: {nested_items}")

    if not isinstance(nested_items, list):
        nested_items = [nested_items]

    nested_processed = []
    for nested_item_index, nested_item in enumerate(nested_items):
        new_indexes = indexes + [nested_item_index]
        nested_item_data = process_each(nested_each_data, data, new_indexes)
        nested_processed.append(nested_item_data)

    return nested_each_data.get('separator', '\n').join(nested_processed)

def process_each(each_data, data, indexes=[]):
    """Process each item in a data structure."""
    key = each_data.get('key')
    items = get_data(data, key, None, indexes)
    print(f"Processing 'each' for key: {key}, Items: {items}")

    if not isinstance(items, list):
        items = [items]

    vars = each_data.get('vars', {})
    all_processed = []

    for item_index, item in enumerate(items):
        new_indexes = indexes + [item_index]
        processed_item = process_item(each_data, data, vars, item, new_indexes)
        all_processed.append(processed_item)

    return each_data.get('separator', '\n').join(all_processed)

def process_item(each_data, data, vars, item, new_indexes):
    """Process an individual item within the 'each' loop."""
    new_content = replace_vars_in_content(each_data.get('content', []), vars, new_indexes)
    formatted_content = process_content(new_content, data, item, new_indexes)
    processed_item = ' '.join(formatted_content)
    print(f"Item at level {len(new_indexes)}: Formatted content: {formatted_content}")
    print(f"Item at level {len(new_indexes)}: Processed item data: {processed_item}")

    # Isolate the nested content processing
    processed_item += process_nested_content(new_content, vars, item, data, new_indexes)

    return processed_item

def process_nested_content(new_content, vars, item, data, new_indexes):
    """Process nested content within an item."""
    nested_processed = []

    for content in new_content:
        if 'each' in content:
            nested_each = content['each'] 
            new_key, _ = replace_key_with_vars(nested_each['key'], vars, new_indexes)
            print(f"new_key: {new_key}")
            nested_each['key'] = new_key
            nested_processed.append(process_nested_items(nested_each, item, data, new_indexes))

    return '\n'.join(nested_processed)

def replace_key_with_vars(key, vars, indexes):
    print(f"Original key: {key}")
    key_parts = key.split('.')
    replaced = False
    for i, part in enumerate(key_parts):
        for var, path in vars.items():
            if var == part:
                try:
                    key_parts[i] = path.format(*indexes)
                    print(f"Replaced '{var}' with '{key_parts[i]}'")
                    replaced = True
                except IndexError as e:
                    print(f"Error: Replacement index out of range for key '{key}': {e}")
                    print(f"Current indexes: {indexes}, Path: {path}")
    new_key = '.'.join(key_parts)
    print(f"New key after replacement: {new_key}")
    return new_key, replaced

def replace_vars_in_content(content, vars, indexes):
    """Replace variable placeholders in content keys with actual indexes."""
    new_content = []
    for content_item in content:
        new_keys = []
        for k in content_item.get('keys', []):
            new_key, _ = replace_key_with_vars(k, vars, indexes)
            new_keys.append(new_key)
        content_item_updated = {**content_item, 'keys': new_keys}
        new_content.append(content_item_updated)
    print(f"new_content: {new_content}")
    return new_content

def set_cell_border(cell, border_name, border_value):
    """Set the border of a cell."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    border = tcBorders.find(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{border_name}')
    if border is None:
        border = OxmlElement(f'w:{border_name}')
        tcBorders.append(border)

    border.set(f'{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}val', border_value)

def hide_table_borders(table):
    """Hide all borders for a table."""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, 'top', 'nil')
            set_cell_border(cell, 'left', 'nil')
            set_cell_border(cell, 'bottom', 'nil')
            set_cell_border(cell, 'right', 'nil')

def apply_cell_style(cell, font_style, font_size):
    """Apply the given font style and size to a table cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if font_style == "bold":
                run.bold = True
            if font_size:  # Assuming font-size is specified in points
                run.font.size = Pt(font_size)

def add_table_section(doc, section, data):
    """Add a table section to the document."""
    tag = section.get('tag')
    align = section.get('align', 'center')
    print(f"Adding table section: {tag}, Align: {align}")

    if 'each' in section:
        items = get_data(data, section['each']['key'])
        vars = section['each']['vars']
        if not isinstance(items, list):
            items = [items]
        
        font_style = section['each'].get('font-style', None)
        font_size = section['each'].get('font-size', None)  # Assuming font-size is specified in points


        for item_index, item in enumerate(items):
            print(f"{tag}: {item}")

            # Get the formats for this item (study or job)
            cell_formats = section['each']['content'][0]['format']

            table = doc.add_table(rows=1, cols=len(cell_formats))
            table.autofit = False
            table.allow_autofit = False
            table.style = 'Table Grid'
            table.alignment = WD_ALIGN_PARAGRAPH.CENTER            

            # Get the page width minus the margins
            page_width = Inches(8.5)  # Adjust this to your page width
            margin = Inches(MARGIN)
            available_width = page_width - (2 * margin)

            # Set the width of each column
            column_widths = section['each'].get('width', [])  # Get column widths from the configuration

            # Calculate the total width used by non-"fitauto" columns
            total_fixed_width = sum(Inches(width) for width in column_widths if width != "fitauto")

            # Calculate the width available for "fitauto" columns
            available_auto_width = available_width - total_fixed_width

            for cell_index, cell_format in enumerate(cell_formats):
                if cell_index < len(column_widths):
                    if column_widths[cell_index] == "fitauto":
                        # Set auto-fit width for the column
                        table.columns[cell_index].width = available_auto_width
                    else:
                        table.columns[cell_index].width = Inches(column_widths[cell_index])  # Set column width
                else:
                    # If not enough widths specified, use a default value (you can adjust as needed)
                    table.columns[cell_index].width = Inches(1.0)

            table.left_margin = Inches(MARGIN)
            table.right_margin = Inches(MARGIN)

            # Initialize an empty list to store cell data
            cell_dataset = []

            indexes = [item_index]

            # Populate the cell data list
            for key in section['each']['content'][0]['keys']:
                print(f"cell_ key1: {key}")
                key, _ = replace_key_with_vars(key, vars, indexes)
                print(f"cell_ key: {key}")
                cell_dataset.append(get_data(data, key))

            print(f"cell_dataset: {cell_dataset}")

            # Define the height for the table rows (in points)
            row_height = section['each'].get('height', 20)
            space_before = section['each'].get('space_before', 2)

            # Populate the table cells with data
            row = table.rows[0]
            row.height = Pt(row_height)
            row.space_before = Pt(space_before)
            for cell_index, cell_format in enumerate(cell_formats):
                cell = row.cells[cell_index]
                cell.text = cell_format.format(*cell_dataset)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Set vertical alignment
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # Set the row height
                cell.height = Pt(row_height)

                # Set the space before for the first paragraph in the cell
                if cell.paragraphs:
                    cell.paragraphs[0].paragraph_format.space_before = Pt(space_before)

                # Apply font style and size
                apply_cell_style(cell, font_style, font_size)

            # Process nested responsibilities if available
            if 'each' in section['each']:
                key, _ = replace_key_with_vars(section['each']['each']['key'], vars, indexes)
                process_nested_responsibilities(table, section['each']['each'], data, key)

                        
            # Hide table borders
            hide_table_borders(table)


def process_nested_responsibilities(table, each_section, data, key):
    responsibilities = get_data(data, key)

    for responsibility in responsibilities:
        row = table.add_row()

        # Set the width of the first cell
        first_cell_width = Inches(1.5)  # Set this as required
        row.cells[0].width = first_cell_width

        # Merge all other cells in the row
        merged_cell = row.cells[1].merge(row.cells[-1])

        formatted_content = process_content(each_section["content"], data, responsibility)
        processed_item = ' '.join(formatted_content)
        merged_cell.text = processed_item
        # Define the height for the table rows (in points)
        row_height = each_section.get('height', [])
        merged_cell.height = Pt(row_height)
        row.height = Pt(row_height)
        merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Set vertical alignment
        merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_section(doc, section, data):
    """Add a section to the document."""
    tag = section.get('tag')
    level = section.get('level', 0)
    content = section.get('content', [])
    align = section.get('align', 'left')
    font_size = section.get('font-size', PARAGRAPH_FONT_SIZE)  # Default to a global font size if not specified
    font_style = section.get('font-style', None)  # Default to a global font size if not specified

    print(f"Adding section: {tag}, Align: {align}, Font size: {font_size}")

    if section['tag'] == 'table':
        add_table_section(doc, section, data)
        return
    elif 'each' in section:
        text = process_each(section['each'], data)
    else:
        text = '\n'.join(process_content(content, data))

    if tag == 'heading':
        paragraph = doc.add_heading(level=level)
        run = paragraph.add_run(text)  # Create a new run for the text
        hex_color = "111111"  # Define the hex color
        rgb_color = hex_to_rgb(hex_color)
        run.font.color.rgb = RGBColor(*rgb_color)
      
        # Adjust spacing for the header based on the level
        if level == 1:
            paragraph.paragraph_format.space_before = Pt(16)  # Adjust as needed
        
        if level == 2:
            run.font.size = Pt(font_size)
    else:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)  # Create a new run for the text
        if font_style == "bold":
            run.font.bold = True
        else:
            run.font.size = Pt(font_size)

    if align == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'right':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def generate_cv_document(data, format_json):
    """Generate a CV document based on the data and format."""
    doc = Document()
    print("Generating CV Document...")
    for section in format_json:
        add_section(doc, section, data)
    return doc

def edit_page_margins(doc):
    # Get the first section of the document (assuming only one section)
    section = doc.sections[0]
    
    # Set left, right, top, and bottom margins (in inches)
    section.left_margin = Inches(LEFT_MARGIN)
    section.right_margin = Inches(RIGHT_MARGIN)
    section.top_margin = Inches(TOP_MARGIN)
    section.bottom_margin = Inches(BOTTOM_MARGIN)

def main():
    if len(sys.argv) > 3:
        json_paths = {}
        for arg in sys.argv[1:-2]:  # Exclude the last two arguments which are format_json_path and output_docx_path
            if '=' in arg:
                key, path = arg.split('=', 1)
                json_paths[key] = path

        format_json_path = sys.argv[-2]
        output_docx_path = sys.argv[-1]

        # Merge JSON data
        cv_data = merge_json_data(json_paths)
    else:
        print("Missing parameters")
        exit(1)
    
    print("Starting CV Generation Script...")
    cv_format = load_json_file(format_json_path)
    cv_document = generate_cv_document(cv_data, cv_format)
    edit_page_margins(cv_document)
    cv_document.save(output_docx_path)
    print("CV Document Generated Successfully.")
    print(f"Opening the CV document: {output_docx_path}")
    open_docx(output_docx_path)

if __name__ == "__main__":
    main()
